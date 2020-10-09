# encoding: utf-8
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Inches
import io
import re
import docx
import requests
tstd = Document()
tstp = tstd.add_paragraph()
tstr = tstp.add_run()
params = {
    'bold':False,
    'italic':False,
    'underline':False,
    'block_quote':False,
    'a_element':False,
    'ul':False,
    'ol':False,
    'image':False
}
def main():
    f = open("bodytest.html","r")
    soup = BeautifulSoup(f, 'lxml')
    print("encoding is {}".format(soup.original_encoding))
    for l in soup.descendants:
        pass
        #print(str(l))
    title = BeautifulSoup('<h1>This is a test</h1>', 'lxml')
    #document = buildDoc(soup, title)
    #document.save('C:\\Users\\etdeh\\Desktop\\test.docx')
    data = {'source':'https://www.zerohedge.com/markets/how-jpms-precious-metals-trading-desk-manipulated-markets-crime-ring-within-bank','title':title,'body':soup,'comments':'!'}
    document = buildZHdoc(**data)
    document.save('C:\\Users\\etdeh\\Desktop\\test.docx')
def buildZHdoc(**kwargs):
    source = kwargs['source']
    title = kwargs['title']
    body = kwargs['body']
    comments = kwargs['comments']
    document = Document()
    if source:
        para = document.add_paragraph()
        para.style = document.styles['Title']
        ttl = title.get_text()
        add_hyperlink(para, ttl, source, False, False, True)
    else:
        document.add_heading(title.get_text(), 0)
    main = body.find_all('div',class_="clearfix")[0].children
    docbody = getBody(main)
    buildDoc(document,docbody)
    return document

def buildDoc(doc,plist):
    pass
    for p in plist:#each p has body child element with list of runs
        ele = p[0]#this is child element
        det = p[1]#this is list of runs with tuple (text, param dictionary)
        pa = doc.add_paragraph()
        for d in det:#each d will be a run tuple
            if d[1] == None:#if no param dictionary, just add element text
                pa.add_run(d[0])
            elif d[1]['block_quote'] or d[1]['ul'] or d[1]['ol']:#test for paragraph level styles
                pass
            elif d[1]['a_element']:#check for image or get href for hyperlink
                pass
            elif d[1]['image']:
                tst = ele.find('img')    
                w = float(tst['width'])/96
                imgdata = requests.get(tst['src']).content
                imgfile = io.BytesIO(imgdata)
                doc.add_picture(imgfile, width=Inches(w))
            #todo  // add code for video and iframe handling
            else:# if no special cases add text with run styles
                r = pa.add_run(d[0])
                r.bold = d[1]['bold']
                r.italic = d[1]['italic']
                r.font.underline = d[1]['underline']

def testName(name):
    if name == 'strong':
        return 'bold'
    elif name == 'em':
        return 'italic'
    elif name == 'a':
        return 'a_element'
    elif name == 'u':
        return 'underline'
    elif name == 'blockquote':
        return 'block_quote'
    elif name == 'ul':
        return 'ul'
    elif name == 'ol':
        return 'ol'
    else:
        return None
    #print("bold is {} and ital is {} and atag is {}".format(bold, ital, atag))
def has_children(element):
    try:
        t=element.contents
    except:
        return False
    else:
        return True
def add_hyperlink(paragraph, text, url, bflag, iflag, head):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)
    r.bold = bflag
    r.italic = iflag
    
    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    if head:
       r.font.underline = False 
    return hyperlink

def getBody(bodyele):
    paragraphs = []
    for child in bodyele:
        #child will be <p>, <ul>, <ol>, <blockquote>, etc... creating line of article
        runs = []
        if child.name == None:#name of None suggests inner text
            runs.append((child,None))
        else:
            #runs = []
            allp = dict(params)
            allp[testName(child.name)] = True
            #print(str(child))
            #c will be <a>, <li>, <bold>, etc... creating runs of paragraph
            for c in child:
                if c.name == None:
                    runs.append((c,allp))
                else:
                    p = dict(allp)#copy params with any paragraph styles
                    p[testName(c.name)] = True #test for run style in case c isn't text
                    runs.extend(testChildren(c,p))
        paragraphs.append((child,runs))
    return paragraphs

def testChildren(element, paramdict):#receive bold element with bold param = true
    lists = []
    print(paramdict)
    for ch in element:
        print(ch.name)
        if ch.name == None:
            lists.append((ch,paramdict))
        elif ch.name == 'picture':
            d = dict(params)
            d['image'] = True
            lists.append((ch,d))
        else:
            d = dict(paramdict)
            d[testName(ch.name)] = True
            if has_children(ch):
                lists.extend(testChildren(ch,d))
        pass
    return lists

if __name__ == "__main__":
    main()
